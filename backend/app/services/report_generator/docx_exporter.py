from __future__ import annotations

import io
from datetime import datetime
from typing import Iterable
from uuid import UUID

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from sqlalchemy.orm import Session

from app.models.project import Project
from app.services.report_generator.aggregator import (
    DocumentBlock,
    ReportContext,
    build_report_context,
)
from app.utils.constants import ReportType


_REPORT_TYPE_LABEL_VI = {
    ReportType.RESEARCH_SUMMARY.value: "Tóm tắt nghiên cứu",
    ReportType.LITERATURE_REVIEW.value: "Tổng quan tài liệu",
    ReportType.DATA_ANALYSIS.value: "Phân tích dữ liệu",
    ReportType.CUSTOM.value: "Báo cáo tùy chỉnh",
}

# Brand color (matches the FE teal palette).
_ACCENT_RGB = RGBColor(0x0F, 0x76, 0x6E)


def _add_paragraph(doc: DocxDocument, text: str | None) -> None:
    if not text:
        return
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)


def _add_bullets(doc: DocxDocument, items: Iterable[str]) -> None:
    items = [i for i in items if i and i.strip()]
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(11)


def _add_kv(doc: DocxDocument, label: str, value: str | None) -> None:
    if not value:
        return
    p = doc.add_paragraph()
    label_run = p.add_run(f"{label}: ")
    label_run.bold = True
    label_run.font.size = Pt(11)
    value_run = p.add_run(value)
    value_run.font.size = Pt(11)


def _add_cover(doc: DocxDocument, ctx: ReportContext) -> None:
    label = _REPORT_TYPE_LABEL_VI.get(ctx.report_type, ctx.report_type)

    # Eyebrow label
    eyebrow = doc.add_paragraph()
    eyebrow.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = eyebrow.add_run(label.upper())
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = _ACCENT_RGB

    # Title
    title = doc.add_heading(ctx.report_title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Subtitle (project topic / description)
    subtitle = ctx.project_topic or ctx.project_description
    if subtitle:
        sub_p = doc.add_paragraph()
        sub_run = sub_p.add_run(subtitle)
        sub_run.italic = True
        sub_run.font.size = Pt(12)

    # Metadata table
    meta_rows = [
        ("Dự án", ctx.project_name or "—"),
        (
            "Tài liệu",
            f"{len(ctx.documents_with_analysis)}/{ctx.total_documents} đã phân tích",
        ),
        ("Tạo lúc", ctx.generated_at.strftime("%d/%m/%Y %H:%M")),
    ]
    table = doc.add_table(rows=len(meta_rows), cols=2)
    table.autofit = True
    for row_idx, (label_, value_) in enumerate(meta_rows):
        row = table.rows[row_idx]
        l_cell, v_cell = row.cells
        for run in l_cell.paragraphs[0].runs:
            run.bold = True
        l_cell.text = label_
        v_cell.text = value_
        for cell in (l_cell, v_cell):
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(11)
        # Bold the labels (text= reset removed bold above; reapply)
        for run in l_cell.paragraphs[0].runs:
            run.bold = True

    doc.add_paragraph()  # spacer


def _add_doc_block(doc: DocxDocument, block: DocumentBlock) -> None:
    doc.add_heading(block.display_title, level=2)

    meta_bits: list[str] = []
    if block.document_type:
        meta_bits.append(f"Loại: {block.document_type}")
    if block.source_url:
        meta_bits.append(f"Nguồn: {block.source_url}")
    if not block.has_analysis:
        meta_bits.append("Chưa có phân tích")
    if meta_bits:
        meta_p = doc.add_paragraph()
        meta_run = meta_p.add_run(" · ".join(meta_bits))
        meta_run.italic = True
        meta_run.font.size = Pt(10)

    if not block.has_analysis:
        _add_paragraph(
            doc,
            "Tài liệu này chưa có phân tích đầy đủ — hãy chạy Analysis "
            "để bổ sung dữ liệu vào báo cáo.",
        )
        return

    if block.summary:
        _add_paragraph(doc, block.summary)

    _add_kv(doc, "Luận điểm chính", block.main_thesis)
    _add_kv(doc, "Đóng góp", block.research_contribution)

    if block.key_findings:
        doc.add_heading("Phát hiện chính", level=3)
        _add_bullets(doc, block.key_findings)

    if block.methodology:
        doc.add_heading("Phương pháp", level=3)
        _add_paragraph(doc, block.methodology)

    if block.limitations:
        doc.add_heading("Giới hạn", level=3)
        _add_bullets(doc, block.limitations)

    if block.keywords:
        _add_kv(doc, "Từ khóa", ", ".join(block.keywords))


def _render_research_summary(doc: DocxDocument, ctx: ReportContext) -> None:
    doc.add_heading("Tóm tắt tổng quan", level=1)
    if ctx.project_topic:
        _add_kv(doc, "Chủ đề", ctx.project_topic)
    if ctx.project_description:
        _add_paragraph(doc, ctx.project_description)
    if ctx.project_research_scope:
        _add_kv(doc, "Phạm vi nghiên cứu", ctx.project_research_scope)
    _add_paragraph(
        doc,
        f"Báo cáo tổng hợp dữ liệu từ {len(ctx.documents_with_analysis)} "
        f"tài liệu đã phân tích trên tổng số {ctx.total_documents} tài liệu "
        "thuộc dự án.",
    )
    if ctx.aggregate_keywords:
        _add_kv(doc, "Từ khóa nổi bật", ", ".join(ctx.aggregate_keywords[:15]))

    doc.add_heading("Phát hiện nổi bật", level=1)
    if ctx.aggregate_findings:
        _add_bullets(doc, ctx.aggregate_findings)
    else:
        _add_paragraph(doc, "Chưa có phát hiện nào được tổng hợp.")

    if ctx.aggregate_research_questions:
        doc.add_heading("Câu hỏi nghiên cứu", level=1)
        _add_bullets(doc, ctx.aggregate_research_questions)

    doc.add_heading("Tài liệu chi tiết", level=1)
    for block in ctx.documents:
        _add_doc_block(doc, block)


def _render_literature_review(doc: DocxDocument, ctx: ReportContext) -> None:
    doc.add_heading("Bối cảnh", level=1)
    _add_kv(doc, "Chủ đề", ctx.project_topic)
    _add_paragraph(doc, ctx.project_description)
    _add_kv(doc, "Phạm vi", ctx.project_research_scope)

    doc.add_heading("Phương pháp tổng hợp", level=1)
    _add_paragraph(
        doc,
        f"Báo cáo tổng hợp dữ liệu từ {len(ctx.documents_with_analysis)}"
        f"/{ctx.total_documents} tài liệu đã được phân tích bằng pipeline "
        "section-grounded của hệ thống. Mỗi tài liệu được trích xuất "
        "theo các trường: tóm tắt, luận điểm chính, phương pháp, đóng góp, "
        "giới hạn, và hướng nghiên cứu tiếp theo.",
    )

    doc.add_heading("Tài liệu tham khảo", level=1)
    for idx, block in enumerate(ctx.documents, 1):
        p = doc.add_paragraph(style="List Number")
        title_run = p.add_run(block.display_title)
        title_run.bold = True
        if block.document_type:
            p.add_run(f"  ({block.document_type})")
        if block.source_url:
            url_run = p.add_run(f"\n   {block.source_url}")
            url_run.font.size = Pt(10)
            url_run.italic = True
        if block.summary:
            sum_p = doc.add_paragraph()
            sum_p.paragraph_format.left_indent = Pt(18)
            sum_run = sum_p.add_run(block.summary)
            sum_run.font.size = Pt(10)

    doc.add_heading("Đóng góp và phát hiện", level=1)
    contribs = [
        (b.display_title, b.research_contribution or b.main_thesis)
        for b in ctx.documents_with_analysis
        if (b.research_contribution or b.main_thesis)
    ]
    if contribs:
        for title, contrib in contribs:
            _add_kv(doc, title, contrib)
    elif ctx.aggregate_findings:
        _add_bullets(doc, ctx.aggregate_findings)

    doc.add_heading("So sánh phương pháp", level=1)
    if ctx.aggregate_methodologies:
        _add_bullets(doc, ctx.aggregate_methodologies)
    else:
        _add_paragraph(doc, "Các tài liệu chưa có thông tin phương pháp.")

    doc.add_heading("Khoảng trống và hướng tiếp theo", level=1)
    doc.add_heading("Giới hạn quan sát được", level=2)
    if ctx.aggregate_limitations:
        _add_bullets(doc, ctx.aggregate_limitations)

    doc.add_heading("Hướng nghiên cứu đề xuất", level=2)
    if ctx.aggregate_future_work:
        _add_bullets(doc, ctx.aggregate_future_work)


def _render_data_analysis(doc: DocxDocument, ctx: ReportContext) -> None:
    doc.add_heading("Phạm vi phân tích", level=1)
    _add_paragraph(
        doc, ctx.project_research_scope or ctx.project_description
    )
    _add_bullets(
        doc,
        [
            f"Tổng số tài liệu: {ctx.total_documents}",
            f"Số tài liệu đã phân tích: {len(ctx.documents_with_analysis)}",
        ],
    )

    doc.add_heading("Bảng tổng quan dữ liệu", level=1)
    if ctx.documents:
        table = doc.add_table(rows=1, cols=6)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        for i, h in enumerate(
            ["#", "Tài liệu", "Loại", "Số phần", "Phát hiện", "Trạng thái"]
        ):
            hdr[i].text = h
            for run in hdr[i].paragraphs[0].runs:
                run.bold = True
        for idx, b in enumerate(ctx.documents, 1):
            row = table.add_row().cells
            row[0].text = str(idx)
            row[1].text = b.display_title
            row[2].text = b.document_type or "—"
            row[3].text = str(b.section_count or "—")
            row[4].text = str(len(b.key_findings) or "—")
            row[5].text = (
                "Đã phân tích" if b.has_analysis else "Chưa phân tích"
            )

    doc.add_heading("Kết quả tổng hợp", level=1)
    doc.add_heading("Phát hiện chính", level=2)
    if ctx.aggregate_findings:
        _add_bullets(doc, ctx.aggregate_findings)
    if ctx.aggregate_keywords:
        doc.add_heading("Từ khóa nổi bật", level=2)
        _add_paragraph(doc, ", ".join(ctx.aggregate_keywords[:20]))
    if ctx.aggregate_limitations:
        doc.add_heading("Cảnh báo về dữ liệu", level=2)
        _add_bullets(doc, ctx.aggregate_limitations)

    doc.add_heading("Phân tích chi tiết theo tài liệu", level=1)
    for block in ctx.documents_with_analysis:
        _add_doc_block(doc, block)


def _render_custom(doc: DocxDocument, ctx: ReportContext) -> None:
    doc.add_heading("Mở đầu", level=1)
    _add_paragraph(
        doc,
        "Đây là báo cáo tùy chỉnh được tạo tự động từ dữ liệu dự án. "
        "Bạn có thể chỉnh sửa nội dung để phù hợp với mục đích trình bày.",
    )
    if ctx.project_description:
        doc.add_heading("Mô tả dự án", level=1)
        _add_paragraph(doc, ctx.project_description)
    doc.add_heading("Tài liệu trong báo cáo", level=1)
    for block in ctx.documents:
        _add_doc_block(doc, block)
    if ctx.aggregate_findings:
        doc.add_heading("Tổng kết phát hiện", level=1)
        _add_bullets(doc, ctx.aggregate_findings)


_DOCX_RENDERERS = {
    ReportType.RESEARCH_SUMMARY.value: _render_research_summary,
    ReportType.LITERATURE_REVIEW.value: _render_literature_review,
    ReportType.DATA_ANALYSIS.value: _render_data_analysis,
    ReportType.CUSTOM.value: _render_custom,
}


def generate_report_docx(
    db: Session,
    project: Project,
    *,
    report_title: str,
    report_type: str,
    included_documents: list[UUID] | None,
) -> bytes:
    """Render the report as a ``.docx`` blob (no temp files)."""
    ctx = build_report_context(
        db=db,
        project=project,
        report_title=report_title,
        report_type=report_type,
        included_documents=included_documents,
    )

    doc = DocxDocument()

    # Tweak default body font to something more legible.
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    _add_cover(doc, ctx)

    renderer = _DOCX_RENDERERS.get(ctx.report_type, _render_custom)
    renderer(doc, ctx)

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run(
        f"Tạo lúc {ctx.generated_at.strftime('%d/%m/%Y %H:%M')} "
        "bằng hệ thống Research Assistant"
    )
    footer_run.italic = True
    footer_run.font.size = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


__all__ = ["generate_report_docx"]
